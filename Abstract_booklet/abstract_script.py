import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_ALIGN_PARAGRAPH  # For text alignment
from docx.shared import Inches  # For setting image size in inches
import re  # For regular expressions to handle multiple spaces
import os

# Load the Excel file
df_plenary = pd.read_excel("Abstracts.xlsx", sheet_name="Plenary")
df_talks = pd.read_excel("Abstracts.xlsx", sheet_name="Talks")
df_posters = pd.read_excel("Abstracts.xlsx", sheet_name="Posters")
df_butlin = pd.read_excel("Abstracts.xlsx", sheet_name="Butlin")

# Rename columns for consistency
df_plenary = df_plenary.rename(columns={
    'what is your talk title?': 'what is your title',
    'What is your talk abstract?': 'What is your abstract'
})

df_talks = df_talks.rename(columns={
    'what is your talk title?': 'what is your title',
    'What is your talk abstract?': 'What is your abstract'
})

df_posters = df_posters.rename(columns={
    'what is your poster title?': 'what is your title',
    'What is your poster abstract?': 'What is your abstract'
})

df_butlin = df_butlin.rename(columns={
    'what is your talk title?': 'what is your title',
    'What is your talk abstract?': 'What is your abstract'
})

# Define the columns to select and clean the abstracts
columns = ['What is your surname?', 'What is your first name?', 'what is your title', 'What is your abstract', 'Please list all the authors with their affiliation below']

def clean_abstract(text):
    """Cleans abstract text by removing unnecessary spaces and newlines, and reducing multiple spaces."""
    if pd.isna(text):  
        return ""
    # Replace multiple spaces with a single space
    cleaned_text = re.sub(r'\s+', ' ', text).strip()  
    return cleaned_text

def clean_affiliation(text):
    """Cleans affiliation text by removing empty lines and reducing multiple spaces."""
    if pd.isna(text):  # Handle missing values gracefully
        return ""
    
    # Split the text into lines
    lines = text.split('\n')
    # Filter out empty lines or lines that only contain whitespace
    cleaned_lines = [line for line in lines if line.strip() != ""]
    # Join the remaining lines back into a single string
    cleaned_affiliation = "\n".join(cleaned_lines)
    # Replace multiple spaces with a single space
    cleaned_affiliation = re.sub(r'\s+', ' ', cleaned_affiliation).strip()

    return cleaned_affiliation

# Clean and prepare each dataframe
def prepare_dataframe(df):
    df_sorted = df.sort_values(by=["What is your surname?", "What is your first name?"], key=lambda col: col.str.lower())
    df_subset = df_sorted[columns].copy()  # No 'Abstract Number' here
    df_subset.loc[:, 'What is your abstract'] = df_subset['What is your abstract'].apply(clean_abstract)
    df_subset.loc[:, 'Please list all the authors with their affiliation below'] = df_subset['Please list all the authors with their affiliation below'].apply(clean_affiliation)
    return df_subset

# Prepare all dataframes
df_plenary_subset = prepare_dataframe(df_plenary)
df_talks_subset = prepare_dataframe(df_talks)
df_posters_subset = prepare_dataframe(df_posters)
df_butlin_subset = prepare_dataframe(df_butlin)
print(df_butlin_subset)

def create_abstract_booklet(df_plenary, df_talks, df_posters, df_butlin):
    # Create a new Document
    doc = Document()
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(0)

    # Add the "Abstracts Booklet" title page
    title_page = doc.add_paragraph()
    title_page_run = title_page.add_run("Population Genetics Group 58\nAbstracts Booklet")
    title_page_run.bold = True
    title_page_run.font.size = Pt(30)  # Larger font for the booklet title
    title_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph()


    # Add an image (for example, a PNG file) and center it
    image_paragraph = doc.add_paragraph()  # Create a new paragraph for the image
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the paragraph
    image_paragraph.add_run().add_picture('Logo.png', width=Inches(5))

    # Add additional logos from the "Logos" folder
    logos_folder = "Logos"  # Folder containing the additional images
    image_files = [f for f in os.listdir(logos_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))]
    
    # Ensure exactly 3 rows (First two rows: 3 images; Last row: 1 centered image)
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add first two rows with 3 images each
    for i in range(6):  # First 6 images (2 rows of 3 images)
        if i % 3 == 0:  # Start a new row every 3 images
            row_cells = table.add_row().cells
    
        # Insert image into the appropriate cell
        image_path = os.path.join(logos_folder, image_files[i])
        paragraph = row_cells[i % 3].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(2.5))
    
    # Add the last row with the single image centered
    last_row = table.add_row().cells
    for i in range(3):  # Merge all three cells in the last row
        if i > 0:
            last_row[0].merge(last_row[i])
    
    # Insert the last image into the centered cell
    paragraph = last_row[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(os.path.join(logos_folder, image_files[6]), width=Inches(2.5))
    doc.add_page_break()

    # Function to add a section
    def add_section(doc, section_title, df, section_type):
        # Add the Section title
        section_paragraph = doc.add_paragraph()
        section_run = section_paragraph.add_run(section_title.upper())
        section_run.bold = True
        section_run.font.size = Pt(36)
        section_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the explanation about numbering
        numbering_paragraph = doc.add_paragraph()
        numbering_paragraph_run = numbering_paragraph.add_run(f"{section_title} are numbered in alphabetical order by author surname")
        numbering_paragraph_run.font.size = Pt(14)
        numbering_paragraph_run.font.name = 'Calibri'
        numbering_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()

        # Initialize the counter
        counter = 1

        # Add abstracts for this section
        for _, row in df.iterrows():
            # Title Numbering (e.g., Plenary 1, Talk 1, etc.)
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(f"{section_type} {counter}: {row['what is your title']}")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.name = 'Calibri'

            # Add a blank line
            doc.add_paragraph()

            # Add Author's Name (underlined and bold)
            author_paragraph = doc.add_paragraph()
            
            author_run = author_paragraph.add_run("Speaker: ")
            author_run.font.bold = True
            author_run.font.size = Pt(15)
            author_run.font.underline = None  # No underline for "Speaker"

            author_run_name = author_paragraph.add_run(f"{row['What is your surname?']}, {row['What is your first name?']}")
            author_run_name.bold = True
            author_run_name.underline = WD_UNDERLINE.SINGLE
            author_run_name.font.size = Pt(15)
            author_run_name.font.name = 'Calibri'

            doc.add_paragraph()

            # Add Author's Affiliations (in a slightly smaller size and not in bold)
            affiliation_paragraph = doc.add_paragraph()
            affiliation_run = affiliation_paragraph.add_run(f"{row['Please list all the authors with their affiliation below']}")
            affiliation_run.font.size = Pt(12)
            affiliation_run.font.name = 'Calibri'

            # Add two blank lines
            doc.add_paragraph()
            doc.add_paragraph()

            # Add the Abstract
            abstract_paragraph = doc.add_paragraph()
            abstract_run = abstract_paragraph.add_run(f"{row['What is your abstract']}")
            abstract_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abstract_run.font.size = Pt(12)

            # Add a page break after each abstract
            doc.add_page_break()

            # Increment the counter
            counter += 1

    # Add Plenary section
    add_section(doc, "Plenary talks", df_plenary, "Plenary")

    # Add Talks section
    add_section(doc, "Talks", df_talks, "Talk")

    # Add Posters section
    add_section(doc, "Posters", df_posters, "Poster")

    # Add Butlin section
    add_section(doc, "Celebration of Roger Butlin's Career Achievements", df_butlin, "Butlin session")

    # Save the document
    doc.save('abstract_booklet.docx')

# Create the booklet
create_abstract_booklet(df_plenary_subset, df_talks_subset, df_posters_subset, df_butlin_subset)
